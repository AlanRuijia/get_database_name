import os
import cx_Oracle
from time import localtime, strftime, sleep
from json import load
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def init():
    os.environ['NLS_LANG'] = 'SIMPLIFIED CHINESE_CHINA.UTF8'
    config = {}
    dir_name = "Data_" + strftime("%Y%m%d_%H%M%S", localtime())
    os.mkdir(dir_name)

    # log = open(dir_name + "/log", "w")
    # log.write(strftime("%Y/%m/%d %H:%M:%S", localtime()) + ": Getting" + dir_name + "\n")

    if os.path.exists("Config.json"):
        fp = open("Config.json", "r")
        config = load(fp)
        fp.close()
    else:
        print "Config.json doesn't exist.\n"
        # log.write("Error " + strftime("%Y/%m/%d %H:%M:%S", localtime()) + ": Config.json doesn't exist.\n")
        sleep(2)
        exit(0)

    last = config["Server"] + "/" + config["SID"]
    con = []
    try:
        con = cx_Oracle.connect(config["UserName"], config["Password"], last)
    except Exception as e:
        print "Error: "
        print e
        exit(0)

    print "Connection Established."
    # log.write(strftime("%Y/%m/%d %H:%M:%S", localtime()) + ": Connection established.\n")

    # expression = '^XIE|^CUX'
    # expression = 'XIE_ERP_GL_INT_T'
    expression = config['expression']

    os.chdir(dir_name)
    return con


def get_table_name(expression, con, shading_elem):
    # Fetch table information
    find_object_info = "select OBJECT_TYPE, OBJECT_NAME, OWNER from ALL_OBJECTS where regexp_like(OBJECT_NAME ,:expression) AND OBJECT_TYPE='TABLE'"
    table_comment_exist = "SELECT COUNT(*) FROM all_tab_comments WHERE table_name = :tableName AND comments IS NOT NULL"
    get_table_comment = "SELECT comment FROM all_tab_comments WHERE table_name = :tableName"
    column_num = "SELECT COUNT(*) FROM all_tab_columns WHERE table_name = :tableName"
    col_com_num = "SELECT COUNT(*) FROM all_col_comments WHERE table_name = :tableName"
    col_comment = "SELECT column_name, comments FROM all_col_comments WHERE table_name = :tableName"

    get_primary_key = '''SELECT cols.column_name
    FROM all_constraints cons, all_cons_columns cols
    WHERE cols.table_name = :tableName
    AND cons.constraint_type = 'P'
    AND cons.constraint_name = cols.constraint_name
    AND cons.owner = cols.owner
    ORDER BY cols.table_name, cols.position'''
    primary_exist = '''SELECT count(constraint_name) FROM all_constraints
                    WHERE table_name = :tableName AND constraint_type = 'P' '''
    get_foreign_key = "SELECT constraint_name FROM all_constraints WHERE table_name = :tableName AND constraint_type = 'R'"
    foreign_exist = '''SELECT count(constraint_name) FROM all_constraints
        WHERE table_name = :tableName AND constraint_type = 'R' '''
    find_creation = "SELECT DBMS_METADATA.GET_DDL(:1,:2,:3) FROM DUAL"

    cur = con.cursor()

    cur.execute(find_object_info, {'expression': str(expression)})
    object_info = cur.fetchall()

    table = Document()
    table_name = Document()
    for res in object_info:
        try:
            # print res
            # Table Creation Query
            cur.execute(table_comment_exist, {'tableName': str(res[1])})
            table_comment_flag = cur.fetchall()

            if not table_comment_flag:
                cur.execute(get_table_comment, {'tableName': str(res[1])})
                table_comment = cur.fetchall()
            else:
                table_comment = ""

            cur.execute(column_num, {'tableName': str(res[1])})
            col_num = cur.fetchall()

            table_name.add_paragraph(res[1] + '   ' + table_comment)
            p_table = table.add_paragraph()
            run = p_table.add_run('Table Name: ' + str(res[1]))
            run.bold = True
            # run.underline = True
            docx_table = table.add_table(rows=1, cols=8, style='TableGrid')
            # docx_table.style = 'TableGrid'
            first_row_cells = docx_table.rows[0].cells
            table_cols = ('Column Name', 'PK', 'Type', 'Size''Scale', 'Null allowed', 'Default', 'Description')
            
            for i in range(0, len(table_cols)):
                # first_row_cells[i].paragraph[0].add_run(table_cols[i]).bold = True
                first_row_cells[i].text = table_cols[i]
            cur.execute(find_creation, res)
            create_info = cur.fetchall()[0][0].read().strip()
            if str(res[1]) == 'AP_CHECK_INTEGERS':
                print "Continue"
                sleep(3)
                continue
            if str(res[1]) == 'FND_DUAL':
                print "Continue"
                sleep(3)
                continue
            print create_info
            print "-----------------------------"
            create_info = create_info.split('\n')
            create_info = create_info[1:col_num[0][0]]
            for i in range(0, len(create_info)):
                create_info[i] = create_info[i].strip()
            print create_info
            create_info[0] = create_info[0][2:]
            print create_info
            print "-----------------------------"

            cur.execute(col_com_num, {'tableName': str(res[1])})
            col_com_flag = cur.fetchall()[0][0]
            col_comments = []
            if col_com_flag:
                cur.execute(col_comment, {'tableName': str(res[1])})
                col_comments = cur.fetchall()

            cur.execute(primary_exist, {'tableName': str(res[1])})
            primary_flag = cur.fetchall()
            primary_name = ""
            if primary_flag[0][0]:
                cur.execute(get_primary_key, {'tableName': str(res[1])})
                primary_name = cur.fetchall()[0][0]
            print "-----------Primary------"
            print primary_flag[0][0]
            print primary_name

            cur.execute(foreign_exist, {'tableName': str(res[1])})
            foreign_flag = cur.fetchall()
            if foreign_flag[0][0]:
                cur.execute(get_foreign_key, {'tableName': str(res[1])})
                foreign_name = cur.fetchall()[0][0]

            for eve in create_info:
                null_allowed = 'Y'
                default_value = ' '
                size = ' '
                scale = ' '
                comment = ' '
                primary = 'N'
                if "NOT NULL ENABLE" in eve:
                    null_allowed = 'N'
                eve = eve.split()
                col_name = eve[0][1:-1]
                if '(' in eve[1]:
                    type_content = eve[1].split('(')[0]
                    type_size = eve[1].split(',')
                    if len(type_size) == 1:
                        size = type_size[0][type_size[0].index('(') + 1:-1]
                    else:
                        size = type_size[0][type_size[0].index('(') + 1:-1]
                        scale = type_size[1][:-1]
                else:
                    type_content = eve[1]
                for i in range(0, len(col_comments)):
                    if col_name == col_comments[i][0]:
                        comment = str(col_comments[i][1])

                if primary_name == col_name:
                    print "Primary Key " + col_name
                    primary = 'Y'

                if "DEFAULT" in eve:
                    default_value = eve[eve.index("DEFAULT") + 1]

                if comment == 'None':
                    comment = ' '
                cells = docx_table.add_row().cells
                cells[0].text = col_name.strip(',')
                cells[1].text = primary.strip(',') 
                cells[2].text = type_content.strip(',') 
                cells[3].text = size.strip(',') 
                cells[4].text = scale.strip(',') 
                cells[5].text = null_allowed.strip(',') 
                cells[6].text = default_value.strip(',')
                cells[7].text = comment.decode('Utf-8')

            table.add_paragraph(' ')
        except IndexError:
            print IndexError
            continue

    table.save('Table.docx')
    table_name.save('TableName.docx')
    cur.close()


def get_view_name(expression, con):
    find_object_info = '''select OBJECT_TYPE, OBJECT_NAME, OWNER from ALL_OBJECTS
                            where regexp_like(OBJECT_NAME ,:expression) and OBJECT_TYPE = 'VIEW' '''
    view_cols = "select column_name from all_tab_columns where table_name = :viewName"
    find_creation = "SELECT DBMS_METADATA.GET_DDL(:1,:2,:3) FROM DUAL"

    cur = con.cursor()
    cur.execute(find_object_info, {'expression': str(expression)})
    object_info = cur.fetchall()

    view = Document()
    for res in object_info:
        view_res = []
        cur.execute(view_cols, {'viewName': str(res[1])})
        view_content = cur.fetchall()
        cur.execute(find_creation, res)
        view_creation_info = cur.fetchall()[0][0].read().strip()
        print view_creation_info
        view_creation_info = view_creation_info.split('\n')
        view_creation_info = view_creation_info[1:]
        for i in range(0, len(view_creation_info)):
            view_creation_info[i].strip()

        view_creation_info[0] = view_creation_info[0][8:]
        print view_content
        print "------------------View Create Info -----------------"
        print view_creation_info
        print "------------------------------len -----------------"
        print len(view_creation_info)

        for i in range(0, min(len(view_content), len(view_creation_info))):
            # view_col_name = view_content[i][0]
            print "------------------i--------------"
            print i
            print len(view_content)
            view_col_name = view_creation_info[i]
            print "------------view_col_name-----------------"
            print view_col_name
            print "------------------split,---specific--------"
            specific_view_col = view_col_name.split(',')
            print specific_view_col
            source_table_flag = 0
            for j in range(0, len(specific_view_col)):
                if str(specific_view_col[j]) == '':
                    continue
                view_col_name = specific_view_col[j].strip().split()
                print "----------view_col_name-------------------"
                print view_col_name
                print "-----------------------------"
                if str(view_col_name[0]).lower() == 'from':
                    source_table_flag = 1
                    view_col_name = view_col_name[1:]
                    source_table = "Source Table|" + str(view_col_name[0]) + " AS " + str(view_col_name[1] + '\n')
                    view_res.insert(0, source_table)
                    continue
                if 'from' in view_col_name:
                    view_col_name = specific_view_col[j].strip().split('from')
                    if len(view_col_name) == 2:
                        view_res.append(view_col_name[0] + '|' + view_col_name[0] + '\n')
                        view_res.insert(0, "Source Table|" + str(view_col_name[1]) + '\n')
                    if len(view_col_name) == 3:
                        view_res.append(str(view_col_name[1]) + '|' + str(view_col_name[0]) + '\n')
                        view_res.insert(0, "Source Table| " + str(view_col_name[2]) + '\n')
                        continue
                if source_table_flag == 1:
                    view_res.insert(0, "Source Table|" + str(view_col_name[0]) + " AS " + str(view_col_name[1]) + '\n')
                    continue
                if len(view_col_name) == 1:
                    view_res.append(str(view_col_name[0]) + '|' + str(view_col_name[0]) + '\n')
                    continue

                view_res.append(str(view_col_name[0]) + '|' + str(view_col_name[1]) + '\n')
            if source_table_flag == 1:
                break

        view_res.insert(1, "View Column|Source Table Column")
        for item in view_res:
            view.write("%s\n" % item)


def get_brief_name(expression, con):
    # Fetch all kinds of object_name from all_objects
    find_object_info = '''select OBJECT_TYPE, OBJECT_NAME, OWNER
                            from ALL_OBJECTS where regexp_like(OBJECT_NAME ,:expression)'''
    cur = con.cursor()

    cur.execute(find_object_info, {'expression': str(expression)})
    object_info = cur.fetchall()

    view_name = Document()
    index_name = Document()
    package_name = Document()
    trigger_name = Document()
    sequence_name = Document()
    for res in object_info:
        # print res
        if str(res[0]) == 'VIEW':
            view_name.add_paragraph(str(res[1]))
        if str(res[0]) == 'INDEX':
            index_name.add_paragraph(str(res[1]))
        if str(res[0]) == 'PACKAGE':
            package_name.add_paragraph(str(res[1]))
        if str(res[0]) == 'TRIGGER':
            trigger_name.add_paragraph(str(res[1]))
        if str(res[0]) == 'SEQUENCE':
            sequence_name.add_paragraph(str(res[1]))
    view_name.save('ViewName.docx')
    index_name.save('IndexName.docx')
    package_name.save('PackageName.docx')
    trigger_name.save('TriggerName.docx')
    sequence_name.save('SequenceName.docx')
con = init()
# Set a cell background (shading) color to RGB A0A0A0(Gray).
shading_elm = parse_xml(r'<w:shd {} w:fill="A0A0A0"/>'.format(nsdecls('w')))
get_table_name('^XIE', con, shading_elm)
get_brief_name('^XIE', con)
